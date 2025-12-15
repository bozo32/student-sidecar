

#!/usr/bin/env python3
"""
extract_text.py — layered document extraction used by build_pairs.py

Goal:
  Provide a single entry point `extract_any(path, prefer_ocr=False, force_academic_pdf=False, grobid_url="http://localhost:8070")`
  that:
    • computes sha256 of the file bytes
    • extracts plain text from PDFs/Word/HTML/TXT
    • optionally calls a local GROBID service to produce TEI for academic PDFs
    • gracefully falls back across multiple extractors
    • optionally OCRs scanned PDFs (via `ocrmypdf --sidecar`)
Returns:
  {
    "sha256": str,
    "text": str or None,
    "tei_xml": str or None,
    "extract_tool": str,   # which extractor produced the text
    "ocr_used": bool,
  }

Notes:
  • We avoid rewriting PDFs with Ghostscript: OCR uses `--sidecar` to capture text only.
  • If `force_academic_pdf` is True and file is .pdf, we try GROBID first; else we try fast text extractors first and GROBID later.
"""

from __future__ import annotations

import io
import os
import re
import sys
import json
import time
import shutil
import hashlib
import logging
import tempfile
import subprocess
from pathlib import Path
from typing import Optional, Dict, Tuple

import requests
import random

# Optional import dependencies. Each is fully optional and failure is handled.
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
except Exception:
    pdfminer_extract_text = None

try:
    from docx import Document as DocxDocument  # python-docx
except Exception:
    DocxDocument = None

try:
    from bs4 import BeautifulSoup  # bs4
except Exception:
    BeautifulSoup = None


log = logging.getLogger(__name__)

# ---- GROBID throttling/retry globals (set in __main__) ----
GROBID_LIMITER = None        # set to RateLimiter instance when --grobid-rps > 0
GROBID_MAX_RETRIES = 3
GROBID_TIMEOUT = 120
OCR_SEMAPHORE = None        # set to threading.Semaphore in __main__ to cap concurrent OCR

import threading

class RateLimiter:
    """
    Simple spacing-based rate limiter: ensures at most `rate_per_sec` calls per second.
    """
    def __init__(self, rate_per_sec: float):
        self.rate = float(rate_per_sec)
        self._lock = threading.Lock()
        self._next_time = 0.0

    def acquire(self):
        if self.rate <= 0:
            return
        # Sleep outside the lock to avoid blocking other threads
        with self._lock:
            now = time.time()
            wait = max(0.0, self._next_time - now)
        if wait > 0:
            time.sleep(wait)
        with self._lock:
            now2 = time.time()
            base = max(now2, self._next_time)
            self._next_time = base + (1.0 / self.rate)


# ---------------------- utilities ----------------------

def sha256_bytes(b: bytes) -> str:
    h = hashlib.sha256()
    h.update(b)
    return h.hexdigest()


def read_bytes(path: Path) -> bytes:
    with path.open("rb") as f:
        return f.read()


def looks_like_pdf(path: Path) -> bool:
    return path.suffix.lower() == ".pdf"


def looks_like_docx(path: Path) -> bool:
    # python-docx supports .docx; legacy .doc is not supported here
    return path.suffix.lower() == ".docx"


def looks_like_html(path: Path) -> bool:
    return path.suffix.lower() in {".html", ".htm"}


def looks_like_text(path: Path) -> bool:
    return path.suffix.lower() in {".txt", ".md", ".rst"}


def _normalize_ws(s: str) -> str:
    # Collapse whitespace but keep paragraph structure roughly intact
    return re.sub(r"[ \t]+", " ", s).replace("\r", "").strip()


# ---------------------- PDF helpers ----------------------

def _pdf_text_with_pymupdf(path: Path) -> Optional[str]:
    if fitz is None:
        return None
    try:
        with fitz.open(str(path)) as doc:
            parts = []
            for page in doc:
                parts.append(page.get_text("text"))
        text = "\n".join(parts)
        text = _normalize_ws(text)
        return text if text else None
    except Exception as e:
        log.debug(f"PyMuPDF failed on {path}: {e}")
        return None


def _pdf_text_with_pdfminer(path: Path, timeout: int = 60) -> Optional[str]:
    # Disabled pdfminer fallback
    # if pdfminer_extract_text is None:
    #     return None
    # try:
    #     # pdfminer can be slow; run in-process but catch errors
    #     text = pdfminer_extract_text(str(path))
    #     text = _normalize_ws(text or "")
    #     return text if text else None
    # except Exception as e:
    #     log.debug(f"pdfminer failed on {path}: {e}")
    #     return None
    return None


def _should_ocr(text: Optional[str], path: Path, allow_long_ocr: bool = False) -> bool:
    """
    Improved heuristic:
      - Only OCR if text_len == 0 and file has >=2 pages.
      - Only OCR if all of the first 3 pages are image-only (no text).
      - If any text is found in first 3 pages, skip OCR.
      - Skip OCR entirely if PDF > 40 pages unless allow_long_ocr is True.
    """
    if fitz is None:
        return False
    try:
        with fitz.open(str(path)) as doc:
            n_pages = doc.page_count
            if n_pages < 2:
                return False
            if n_pages > 40 and not allow_long_ocr:
                return False
            # Only OCR if text_len == 0
            if text and len(text.strip()) >= 200:
                return False
            # Check first up to 3 pages for any text
            max_pages_to_check = min(3, n_pages)
            for i in range(max_pages_to_check):
                pg = doc.load_page(i)
                txt = pg.get_text("text")
                if txt and txt.strip():
                    return False
            # All checked pages are image-only
            return True
    except Exception as e:
        log.debug(f"_should_ocr failed: {e}")
        return False


def _available(cmd: str) -> bool:
    return shutil.which(cmd) is not None


def _ocr_text_sidecar(path: Path, timeout: int = 300) -> Optional[str]:
    """Run OCR via ocrmypdf and return extracted text (sidecar only).

    We use `ocrmypdf --skip-text --sidecar out.txt in.pdf out.pdf`.
    If OCR_SEMAPHORE is configured, it gates concurrent OCR jobs.
    """
    if not _available("ocrmypdf"):
        return None

    sema = OCR_SEMAPHORE
    if sema is None:
        return _ocr_text_sidecar_impl(path, timeout=timeout)

    with sema:
        return _ocr_text_sidecar_impl(path, timeout=timeout)


# Implementation for _ocr_text_sidecar (semaphore wrapper calls this).
def _ocr_text_sidecar_impl(path: Path, timeout: int = 300) -> Optional[str]:
    """Implementation for _ocr_text_sidecar (semaphore wrapper calls this)."""
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        out_pdf = td / "out.pdf"
        sidecar = td / "out.txt"
        cmd = [
            "ocrmypdf",
            "--skip-text",              # don't overwrite existing text objects
            "--sidecar", str(sidecar),  # write text here
            "--quiet",
            str(path),
            str(out_pdf),
        ]
        try:
            proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)
            if proc.returncode != 0:
                log.debug(f"ocrmypdf failed: rc={proc.returncode} stderr={proc.stderr.decode(errors='ignore')[:500]}")
                return None
            if sidecar.exists():
                txt = sidecar.read_text(encoding="utf-8", errors="ignore")
                txt = _normalize_ws(txt)
                return txt if txt else None
        except subprocess.TimeoutExpired:
            log.debug("ocrmypdf timed out")
        except Exception as e:
            log.debug(f"ocrmypdf error: {e}")
    return None




# ---------------------- GROBID ----------------------

def _grobid_fulltext_tei(path: Path, grobid_url: str = "http://localhost:8070", timeout: Optional[int] = None, metrics: Optional[dict] = None) -> Optional[str]:
    """
    Call a local GROBID service to get TEI XML for the PDF.
    - Respects a global RateLimiter (GROBID_LIMITER)
    - Retries on 429/5xx with exponential backoff + jitter
    Endpoint: POST {grobid_url}/api/processFulltextDocument
    """
    url = f"{grobid_url.rstrip('/')}/api/processFulltextDocument"
    data = {
        "consolidateHeader": "0",
        "consolidateCitations": "0",
    }
    max_retries = int(GROBID_MAX_RETRIES)
    eff_timeout = int(timeout if timeout is not None else GROBID_TIMEOUT)

    m = metrics if isinstance(metrics, dict) else None
    grobid_request_s = 0.0
    grobid_retry_sleep_s = 0.0
    grobid_attempts = 0
    grobid_last_status = None

    for attempt in range(max_retries + 1):
        # client-side rate limit
        if GROBID_LIMITER is not None:
            try:
                GROBID_LIMITER.acquire()
            except Exception:
                pass

        try:
            t_req0 = time.time()
            with path.open("rb") as fh:
                files = {"input": (path.name, fh, "application/pdf")}
                r = requests.post(url, files=files, data=data, timeout=eff_timeout)
            t_req1 = time.time()
            grobid_request_s += (t_req1 - t_req0)
            grobid_attempts += 1
            grobid_last_status = r.status_code
            # Handle rate-limit / server overload explicitly
            if r.status_code in (429, 503, 502, 500):
                # backoff with jitter
                if attempt < max_retries:
                    sleep_s = min(30.0, (2 ** attempt)) + random.random()
                    log.debug(f"GROBID {r.status_code} on {path.name}; retrying in {sleep_s:.2f}s (attempt {attempt+1}/{max_retries})")
                    time.sleep(sleep_s)
                    grobid_retry_sleep_s += sleep_s
                    continue
                else:
                    log.debug(f"GROBID gave status {r.status_code} after retries for {path.name}")
                    if m is not None:
                        m["grobid_request_s"] = float(grobid_request_s)
                        m["grobid_retry_sleep_s"] = float(grobid_retry_sleep_s)
                        m["grobid_attempts"] = int(grobid_attempts)
                        m["grobid_last_status"] = grobid_last_status
                        m["grobid_total_s"] = float(grobid_request_s + grobid_retry_sleep_s)
                    return None

            r.raise_for_status()
            tei = r.text
            if tei and ("<TEI" in tei or "<tei" in tei or tei.lstrip().startswith("<")):
                if m is not None:
                    m["grobid_request_s"] = float(grobid_request_s)
                    m["grobid_retry_sleep_s"] = float(grobid_retry_sleep_s)
                    m["grobid_attempts"] = int(grobid_attempts)
                    m["grobid_last_status"] = grobid_last_status
                    m["grobid_total_s"] = float(grobid_request_s + grobid_retry_sleep_s)
                return tei
            # Non-TEI content: no retry unless we think it's a transient HTTP error (already handled)
            if m is not None:
                m["grobid_request_s"] = float(grobid_request_s)
                m["grobid_retry_sleep_s"] = float(grobid_retry_sleep_s)
                m["grobid_attempts"] = int(grobid_attempts)
                m["grobid_last_status"] = grobid_last_status
                m["grobid_total_s"] = float(grobid_request_s + grobid_retry_sleep_s)
            return None

        except requests.Timeout:
            grobid_attempts += 1
            grobid_last_status = "timeout"
            if attempt < max_retries:
                sleep_s = min(30.0, (2 ** attempt)) + random.random()
                log.debug(f"GROBID timeout on {path.name}; retrying in {sleep_s:.2f}s (attempt {attempt+1}/{max_retries})")
                time.sleep(sleep_s)
                grobid_retry_sleep_s += sleep_s
                continue
            else:
                log.debug(f"GROBID timeout after retries on {path.name}")
                if m is not None:
                    m["grobid_request_s"] = float(grobid_request_s)
                    m["grobid_retry_sleep_s"] = float(grobid_retry_sleep_s)
                    m["grobid_attempts"] = int(grobid_attempts)
                    m["grobid_last_status"] = grobid_last_status
                    m["grobid_total_s"] = float(grobid_request_s + grobid_retry_sleep_s)
                return None
        except Exception as e:
            grobid_attempts += 1
            grobid_last_status = f"error:{type(e).__name__}"
            # Only retry for network-ish errors; otherwise, bail
            if attempt < max_retries:
                sleep_s = min(15.0, 0.5 * (2 ** attempt)) + random.random()
                log.debug(f"GROBID error on {path.name}: {e}; retrying in {sleep_s:.2f}s (attempt {attempt+1}/{max_retries})")
                time.sleep(sleep_s)
                grobid_retry_sleep_s += sleep_s
                continue
            else:
                log.debug(f"GROBID error after retries on {path.name}: {e}")
                if m is not None:
                    m["grobid_request_s"] = float(grobid_request_s)
                    m["grobid_retry_sleep_s"] = float(grobid_retry_sleep_s)
                    m["grobid_attempts"] = int(grobid_attempts)
                    m["grobid_last_status"] = grobid_last_status
                    m["grobid_total_s"] = float(grobid_request_s + grobid_retry_sleep_s)
                return None

    if m is not None:
        m["grobid_request_s"] = float(grobid_request_s)
        m["grobid_retry_sleep_s"] = float(grobid_retry_sleep_s)
        m["grobid_attempts"] = int(grobid_attempts)
        m["grobid_last_status"] = grobid_last_status
        m["grobid_total_s"] = float(grobid_request_s + grobid_retry_sleep_s)
    return None


# ---------------------- Other file types ----------------------

def _docx_text(path: Path) -> Optional[str]:
    if DocxDocument is None:
        return None
    try:
        doc = DocxDocument(str(path))
        parts = []
        for p in doc.paragraphs:
            parts.append(p.text)
        # tables
        for table in doc.tables:
            for row in table.rows:
                parts.append("  ".join(cell.text for cell in row.cells))
        text = "\n".join(parts)
        text = _normalize_ws(text)
        return text if text else None
    except Exception as e:
        log.debug(f"python-docx failed: {e}")
        return None


def _html_text(path: Path) -> Optional[str]:
    if BeautifulSoup is None:
        return None
    try:
        html = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")
        # Drop nav/script/style
        for tag in soup(["script", "style", "nav", "header", "footer"]):
            tag.extract()
        txt = soup.get_text("\n")
        txt = _normalize_ws(txt)
        return txt if txt else None
    except Exception as e:
        log.debug(f"BeautifulSoup failed: {e}")
        return None


def _txt_text(path: Path) -> Optional[str]:
    try:
        txt = path.read_text(encoding="utf-8", errors="ignore")
        txt = _normalize_ws(txt)
        return txt if txt else None
    except Exception as e:
        log.debug(f"read txt failed: {e}")
        return None


# ---------------------- main entry ----------------------

def extract_any(path: Path,
                prefer_ocr: bool = False,
                force_academic_pdf: bool = False,
                grobid_url: str = "http://localhost:8070",
                allow_long_ocr: bool = False,
                timing_log: Optional[dict] = None,
                ) -> Dict[str, object]:
    """
    Unified extractor.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(str(path))

    byts = read_bytes(path)
    sha = sha256_bytes(byts)

    tei_xml: Optional[str] = None
    text: Optional[str] = None
    tool: Optional[str] = None
    ocr_used = False
    timings = timing_log if timing_log is not None else {}

    # PDF path
    if looks_like_pdf(path):
        # Strategy ordering
        tried = []

        def try_grobid() -> Optional[str]:
            grobid_metrics: dict = {}
            tei = _grobid_fulltext_tei(path, grobid_url=grobid_url, timeout=GROBID_TIMEOUT, metrics=grobid_metrics)
            # Copy grobid metrics into timings with explicit keys
            for k, v in grobid_metrics.items():
                timings[k] = v
            nonlocal tei_xml
            if tei:
                tei_xml = tei
                return tei
            return None

        def try_pymupdf() -> Optional[str]:
            t0 = time.time()
            res = _pdf_text_with_pymupdf(path)
            t1 = time.time()
            timings["pymupdf_s"] = t1 - t0
            return res

        def try_pdfminer() -> Optional[str]:
            t0 = time.time()
            res = _pdf_text_with_pdfminer(path)
            t1 = time.time()
            timings["pdfminer"] = t1 - t0
            return res

        def try_ocr() -> Optional[str]:
            t0 = time.time()
            res = _ocr_text_sidecar(path)
            t1 = time.time()
            timings["ocr_s"] = t1 - t0
            return res

        if force_academic_pdf:
            tei = try_grobid()
            # Regardless of TEI success, attempt fast text extraction
            text = try_pymupdf() # or try_pdfminer()
            if tei and text:
                tool = "grobid+pymupdf"
            elif tei:
                tool = "grobid"
            elif text:
                tool = "pymupdf"
            # OCR heuristic
            if _should_ocr(text, path, allow_long_ocr=allow_long_ocr) and prefer_ocr:
                ocr_txt = try_ocr()
                if ocr_txt and (not text or len(ocr_txt) > len(text)):
                    text = ocr_txt
                    tool = (tool + "+ocr") if tool else "ocrmypdf"
                    ocr_used = True
        else:
            # Try fast text first
            text = try_pymupdf()
            tool = "pymupdf" if text else None
            # pdfminer fallback disabled
            # if not text:
            #     text = try_pdfminer()
            #     tool = "pdfminer" if text else None

            if _should_ocr(text, path, allow_long_ocr=allow_long_ocr) and prefer_ocr:
                ocr_txt = try_ocr()
                if ocr_txt and (not text or len(ocr_txt) > len(text)):
                    text = ocr_txt
                    tool = "ocrmypdf"
                    ocr_used = True

            # Optionally also get TEI if it looks academic or if forced
            # lightweight heuristic: if file name contains conf/journal-ish words
            if ("doi" in path.name.lower() or "conf" in path.name.lower() or "journal" in path.name.lower()) or force_academic_pdf:
                tei = try_grobid()
                if tei:
                    tei_xml = tei
                    tool = (tool + "+grobid") if tool else "grobid"

        return {
            "sha256": sha,
            "text": text,
            "tei_xml": tei_xml,
            "extract_tool": tool or "unknown",
            "ocr_used": bool(ocr_used),
            "timings": timings,
        }

    # DOCX
    if looks_like_docx(path):
        t0 = time.time()
        text = _docx_text(path)
        t1 = time.time()
        timings["docx_s"] = t1 - t0
        return {
            "sha256": sha,
            "text": text,
            "tei_xml": None,
            "extract_tool": "python-docx" if text else "unknown",
            "ocr_used": False,
            "timings": timings,
        }

    # HTML
    if looks_like_html(path):
        t0 = time.time()
        text = _html_text(path)
        t1 = time.time()
        timings["html_s"] = t1 - t0
        return {
            "sha256": sha,
            "text": text,
            "tei_xml": None,
            "extract_tool": "beautifulsoup" if text else "unknown",
            "ocr_used": False,
            "timings": timings,
        }

    # TXT
    if looks_like_text(path):
        t0 = time.time()
        text = _txt_text(path)
        t1 = time.time()
        timings["txt_s"] = t1 - t0
        return {
            "sha256": sha,
            "text": text,
            "tei_xml": None,
            "extract_tool": "plaintext" if text else "unknown",
            "ocr_used": False,
            "timings": timings,
        }

    # Fallback: unknown binary → no text

    return {
        "sha256": sha,
        "text": None,
        "tei_xml": None,
        "extract_tool": "unsupported",
        "ocr_used": False,
        "timings": timings,
    }


# ---------------------- multiprocessing-safe worker ----------------------
# NOTE: On macOS the default multiprocessing start method is `spawn`, which requires
# worker callables to be importable at module top-level (picklable). Do not move this
# inside __main__.

def extract_other_worker(fpath_str: str,
                         allow_long_ocr: bool,
                         grobid_url: str) -> tuple[str, Optional[Dict[str, object]], Optional[Exception]]:
    """ProcessPool worker for non-PDF files (and PDFs when we intentionally avoid GROBID).

    Returns: (fpath_str, out_dict_or_none, err_or_none)
    """
    timing_log: dict = {}
    try:
        t0 = time.time()
        out = extract_any(
            Path(fpath_str),
            prefer_ocr=False,
            force_academic_pdf=False,
            grobid_url=grobid_url,
            allow_long_ocr=bool(allow_long_ocr),
            timing_log=timing_log,
        )
        t1 = time.time()
        out.setdefault("timings", {})
        out["timings"]["wall_elapsed"] = t1 - t0
        return (fpath_str, out, None)
    except Exception as e:
        return (fpath_str, None, e)

# ---------------------- CLI runner ----------------------
if __name__ == "__main__":
    import argparse
    from collections import defaultdict
    import pandas as pd
    import concurrent.futures
    import threading
    import hashlib

    parser = argparse.ArgumentParser(description="Extract text/TEI from a tree of documents and write a sources.parquet manifest.")
    parser.add_argument("root", type=str, help="Root directory to walk (group folders under here).")
    parser.add_argument("--texts-dir", type=str, default="artifacts/text", help="Directory to write .txt and .tei.xml artifacts.")
    parser.add_argument("--parquet-dir", type=str, default="artifacts/parquet", help="Directory to write sources.parquet.")
    parser.add_argument("--out-root", type=str, default=None, help="Cohort output root; when set, parquet goes to <out-root>/parquet.")
    parser.add_argument("--store-dir", type=str, default=None, help="Global store root; when set, artifacts go to <store-dir>/text and <store-dir>/meta.")
    parser.add_argument("--cohort-id", type=str, default=None, help="Cohort identifier recorded in sources.parquet (default derives from out-root name).")
    parser.add_argument("--cohort-note", type=str, default=None,
                        help="Short human-readable description of the cohort/class; stored in sources.parquet.")
    parser.add_argument("--grobid-url", type=str, default="http://localhost:8070", help="Base URL of the GROBID service.")
    parser.add_argument("--force-academic-pdf", action="store_true", help="Try GROBID first for PDFs (assume academic PDFs).")
    parser.add_argument("--prefer-ocr", action="store_true", help="If PDF text looks sparse, allow OCR sidecar and prefer it when longer.")
    parser.add_argument("--reextract", action="store_true", help="Overwrite existing text/tei artifacts; otherwise skip files already extracted.")
    parser.add_argument("--no-merge-manifest", action="store_true",
                        help="Do not merge with an existing sources.parquet; overwrite it with only this run's records.")
    parser.add_argument("--extensions", type=str, default=".pdf,.docx,.html,.htm,.txt,.md,.rst", help="Comma-separated list of file extensions to process.")
    parser.add_argument("--verify", action="store_true", help="After (or without) extraction, print a per-group verification report.")
    parser.add_argument("--verify-only", action="store_true", help="Do not extract; only build/print a verification report against existing artifacts/parquet.")
    parser.add_argument("--verify-report", type=str, default=None, help="Optional path to write a JSON verification report.")
    parser.add_argument("--workers", type=int, default=1, help="Number of concurrent extraction workers (for GROBID/threaded requests).")
    parser.add_argument("--allow-long-ocr", action="store_true", help="Allow OCR on PDFs with more than 40 pages.")
    parser.add_argument("--grobid-rps", type=float, default=0.5, help="Client-side throttle for GROBID: max requests per second (0 = no throttle).")
    parser.add_argument("--grobid-max-retries", type=int, default=3, help="Max retries for GROBID 429/5xx/timeouts.")
    parser.add_argument("--grobid-timeout", type=int, default=60, help="Per-request timeout (seconds) for GROBID.")
    args = parser.parse_args()

    root = Path(args.root).expanduser().resolve()

    # Cohort-scoped output root (parquet/verification/etc.)
    out_root = Path(args.out_root).expanduser().resolve() if args.out_root else None
    if out_root is not None:
        parquet_dir = out_root / "parquet"
    else:
        parquet_dir = Path(args.parquet_dir).expanduser().resolve()
    parquet_dir.mkdir(parents=True, exist_ok=True)

    # Global, content-addressed store for extracted artifacts (text/tei)
    store_dir = Path(args.store_dir).expanduser().resolve() if args.store_dir else None
    if store_dir is not None:
        texts_dir = store_dir / "text"
        meta_dir = store_dir / "meta"
    else:
        texts_dir = Path(args.texts_dir).expanduser().resolve()
        meta_dir = None

    texts_dir.mkdir(parents=True, exist_ok=True)
    if meta_dir is not None:
        meta_dir.mkdir(parents=True, exist_ok=True)

    # Cohort identifier (used in sources.parquet)
    if args.cohort_id:
        cohort_id = str(args.cohort_id)
    elif out_root is not None:
        cohort_id = out_root.name
    else:
        from datetime import datetime, timezone
        cohort_id = datetime.now().astimezone().strftime("%Y%m%d-%H%M%S%z")
    cohort_note = args.cohort_note

    # Allow passing a single file path as `root`.
    root_is_file = root.is_file()
    if root_is_file:
        single_file = root
        root_dir_for_group = root.parent
    else:
        single_file = None
        root_dir_for_group = root

    # Initialize GROBID throttling/retry globals from CLI
    GROBID_MAX_RETRIES = int(args.grobid_max_retries)
    GROBID_TIMEOUT = int(args.grobid_timeout)
    GROBID_LIMITER = RateLimiter(args.grobid_rps) if args.grobid_rps and args.grobid_rps > 0 else None

    wanted_exts = {e.strip().lower() if e.strip().startswith(".") else f".{e.strip().lower()}"
                   for e in args.extensions.split(",") if e.strip()}

    def group_from_path(p: Path, root_dir: Path) -> str:
        try:
            rel = p.resolve().relative_to(root_dir)
            parts = rel.parts
            return parts[0] if parts else "_root"
        except Exception:
            return "_unknown"

    records = []
    seen_sha = set()

    def run_verification(expected_paths: set[str], df_or_none):
        """
        expected_paths: set of all eligible source file paths discovered under root (filtered by extension).
        df_or_none: either a pandas DataFrame with records OR None (we may try to load parquet if present).

        Strategy:
          1) Always compute artifact-based verification for every expected path (by hashing the file and
             checking for {sha}.txt or {sha}.tei.xml in texts_dir). This captures work done even if the
             manifest wasn't written.
          2) If a parquet manifest exists (or df provided), merge it on top (logical OR) so any positive
             `extract_ok` marks paths as extracted even if artifacts moved, and also handles deduped
             files that appear under multiple paths.
        """
        import glob
        df = df_or_none
        out_path = parquet_dir / "sources.parquet"
        if df is None and out_path.exists():
            try:
                df = pd.read_parquet(out_path)
            except Exception as e:
                print(f"[VERIFY] Could not load {out_path}: {e}")
                df = None

        # 1) Artifact-based verification for ALL expected paths
        path_ok: dict[str, bool] = {}
        for src_path in expected_paths:
            try:
                with open(src_path, "rb") as fh:
                    b = fh.read()
                sha = hashlib.sha256(b).hexdigest()
                txt_file = texts_dir / f"{sha}.txt"
                tei_file = texts_dir / f"{sha}.tei.xml"
                ok = txt_file.exists() or tei_file.exists()
                path_ok[src_path] = ok
            except Exception:
                path_ok[src_path] = False

        # 2) Merge any parquet info (logical OR)
        if df is not None and not df.empty:
            # Normalize any array-likes in all_paths, and OR in extract_ok
            for _, row in df.iterrows():
                ok = bool(row.get("extract_ok"))
                paths = row.get("all_paths")
                if paths is None:
                    paths = []
                elif isinstance(paths, str):
                    paths = [paths]
                elif hasattr(paths, "tolist"):  # numpy array
                    paths = paths.tolist()
                for p in paths:
                    sp = str(p)
                    path_ok[sp] = path_ok.get(sp, False) or ok

        def grp(p: str) -> str:
            return group_from_path(Path(p), root_dir_for_group)

        groups = {}
        for p in expected_paths:
            g = grp(p)
            info = groups.setdefault(g, {"expected": 0, "ok": 0, "missing": []})
            info["expected"] += 1
            if path_ok.get(p, False):
                info["ok"] += 1
            else:
                info["missing"].append(p)

        print("\n[VERIFY] Per-group extraction status")
        for g, info in sorted(groups.items()):
            exp = info["expected"]
            ok = info["ok"]
            miss = exp - ok
            status = "✅" if miss == 0 else ("⚠️" if miss <= 3 else "❌")
            print(f"[VERIFY] {g}: {ok}/{exp} extracted {status}")
            if miss > 0:
                for p in info["missing"][:10]:
                    print(f"         missing: {p}")
                if miss > 10:
                    print(f"         ... and {miss-10} more")

        if args.verify_report:
            try:
                report = {g: {"expected": v["expected"], "ok": v["ok"], "missing": v["missing"]} for g, v in groups.items()}
                Path(args.verify_report).parent.mkdir(parents=True, exist_ok=True)
                with open(args.verify_report, "w", encoding="utf-8") as fh:
                    json.dump(report, fh, indent=2)
                print(f"[VERIFY] Wrote JSON report to {args.verify_report}")
            except Exception as e:
                print(f"[VERIFY] Could not write report: {e}")

    def write_artifacts(src_path: Path, out: Dict[str, object]) -> tuple[Optional[str], Optional[str]]:
        """Write content-addressed sidecar artifacts.

        Returns (text_path_str_or_none, tei_path_str_or_none) for what was actually written/present.
        """
        sha = (out.get("sha256") or "").strip().lower()
        if not sha:
            return (None, None)

        txt_path = texts_dir / f"{sha}.txt"
        tei_path = texts_dir / f"{sha}.tei.xml"

        text = out.get("text")
        tei_xml = out.get("tei_xml")

        wrote_txt = False
        wrote_tei = False

        # If not reextract, only write missing artifacts; if reextract, overwrite.
        try:
            if isinstance(text, str) and text.strip():
                if args.reextract or (not txt_path.exists()):
                    txt_path.write_text(text, encoding="utf-8", errors="ignore")
                    wrote_txt = True
            if isinstance(tei_xml, str) and tei_xml.strip():
                if args.reextract or (not tei_path.exists()):
                    tei_path.write_text(tei_xml, encoding="utf-8", errors="ignore")
                    wrote_tei = True
        except Exception as e:
            # Record the failure but do not crash the whole run.
            log.debug(f"write_artifacts failed for sha={sha} path={src_path}: {e}")

        # Optional: write per-source provenance metadata into the global store
        if meta_dir is not None:
            try:
                meta_path = meta_dir / f"{sha}.json"
                payload = {
                    "sha256": sha,
                    "first_seen_path": str(src_path),
                    "extract_tool": out.get("extract_tool"),
                    "ocr_used": bool(out.get("ocr_used")),
                    "timings": out.get("timings", {}),
                    "grobid_url": args.grobid_url if hasattr(args, "grobid_url") else None,
                    "cohort_id": cohort_id,
                    "cohort_note": cohort_note,
                }
                meta_path.write_text(
                    json.dumps(payload, ensure_ascii=False, indent=2),
                    encoding="utf-8",
                    errors="ignore",
                )
            except Exception:
                pass

        # Return what exists on disk now.
        txt_out = str(txt_path) if txt_path.exists() else None
        tei_out = str(tei_path) if tei_path.exists() else None

        # If we expected to write something but it still doesn't exist, warn loudly.
        if isinstance(text, str) and text.strip() and txt_out is None:
            print(f"[WARN] expected text sidecar missing after write: {txt_path}")
        if isinstance(tei_xml, str) and tei_xml.strip() and tei_out is None:
            print(f"[WARN] expected TEI sidecar missing after write: {tei_path}")

        return (txt_out, tei_out)
    # For verification we collect all eligible source paths, even when not extracting.
    all_expected_paths: set[str] = set()
    files_to_extract: list[Path] = []

    if root_is_file:
        fpath = single_file
        if fpath is not None and fpath.exists() and not fpath.name.startswith("."):
            if fpath.suffix.lower() in wanted_exts:
                all_expected_paths.add(str(fpath))
                if not args.verify_only:
                    files_to_extract.append(fpath)
    else:
        for dirpath, dirnames, filenames in os.walk(root):
            dirnames[:] = [d for d in dirnames if not d.startswith(".")]
            for fname in filenames:
                if fname.startswith("."):
                    continue
                fpath = Path(dirpath) / fname
                if fpath.suffix.lower() in wanted_exts:
                    all_expected_paths.add(str(fpath))
                if fpath.suffix.lower() not in wanted_exts:
                    continue
                if args.verify_only:
                    continue
                files_to_extract.append(fpath)

    # Concurrency
    max_workers = max(1, args.workers)
    max_ocr_workers = 2
    OCR_SEMAPHORE = threading.Semaphore(max_ocr_workers)
    # Partition files by extension for GROBID (PDFs) vs local (others)
    pdf_files = [p for p in files_to_extract if p.suffix.lower() == ".pdf"]
    other_files = [p for p in files_to_extract if p.suffix.lower() != ".pdf"]

    # Extraction function for PDFs (may call GROBID, OCR, etc)
    def extract_pdf_wrapper(fpath):
        timing_log = {}
        try:
            t0 = time.time()
            out = extract_any(
                fpath,
                prefer_ocr=bool(args.prefer_ocr),
                force_academic_pdf=bool(args.force_academic_pdf),
                grobid_url=args.grobid_url,
                allow_long_ocr=bool(args.allow_long_ocr),
                timing_log=timing_log,
            )
            t1 = time.time()
            out["timings"]["wall_elapsed"] = t1 - t0
            return (fpath, out, None)
        except Exception as e:
            return (fpath, None, e)


    # Use ThreadPoolExecutor for GROBID/PDFs, ProcessPoolExecutor for local
    pdf_results = []
    other_results = []
    if pdf_files:
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as pool:
            futs = [pool.submit(extract_pdf_wrapper, p) for p in pdf_files]
            for fut in concurrent.futures.as_completed(futs):
                pdf_results.append(fut.result())
    if other_files:
        # On macOS, prefer `fork` to reduce pickling/purity issues for large local extraction.
        # If `fork` is unavailable (or on other OSes), fall back to the default context.
        import multiprocessing as mp
        mp_ctx = mp.get_context("fork") if sys.platform == "darwin" else None

        with concurrent.futures.ProcessPoolExecutor(max_workers=max_workers, mp_context=mp_ctx) as pool:
            futs = [
                pool.submit(extract_other_worker, str(p), bool(args.allow_long_ocr), str(args.grobid_url))
                for p in other_files
            ]
            for fut in concurrent.futures.as_completed(futs):
                other_results.append(fut.result())

    # Combine results in original order
    fpath_to_result = {}
    for fpath, out, err in pdf_results + other_results:
        # pdf_results returns Path; process pool returns string
        key = str(fpath)
        fpath_to_result[key] = (out, err)

    for fpath in files_to_extract:
        out, err = fpath_to_result.get(str(fpath), (None, None))
        if err:
            records.append({
                "source_sha256": None,
                "canonical_ext": fpath.suffix.lower().lstrip("."),
                "bytes": fpath.stat().st_size if fpath.exists() else None,
                "mime": None,
                "first_seen_path": str(fpath),
                "all_paths": [str(fpath)],
                "text_path": None,
                "tei_path": None,
                "extract_tool": f"error:{type(err).__name__}",
                "ocr_used": None,
                "extract_ok": False,
                "extract_error": str(err),
                "cohort_id": cohort_id,
                "cohort_note": cohort_note,
                "group_id": group_from_path(fpath, root_dir_for_group),
                "group": group_from_path(fpath, root_dir_for_group),
            })
            print(f"[EXTRACT-ERROR] {fpath}: {err}")
            continue
        if out is None:
            print(f"[WARN] No output for {fpath}")
            continue
        sha = out.get("sha256")
        if not sha:
            print(f"[WARN] No sha produced for {fpath}")
            continue
        actual_txt_path, actual_tei_path = write_artifacts(fpath, out)
        rec = {
            "source_sha256": sha,
            "canonical_ext": fpath.suffix.lower().lstrip("."),
            "bytes": fpath.stat().st_size,
            "mime": None,
            "first_seen_path": str(fpath),
            "all_paths": [str(fpath)],
            "text_path": actual_txt_path,
            "tei_path": actual_tei_path,
            "extract_tool": out.get("extract_tool"),
            "ocr_used": bool(out.get("ocr_used")),
            "extract_ok": bool(actual_txt_path or actual_tei_path),
            "extract_error": None,
            "cohort_id": cohort_id,
            "cohort_note": cohort_note,
            "group_id": group_from_path(fpath, root_dir_for_group),
            "group": group_from_path(fpath, root_dir_for_group),
            "timings": out.get("timings", {}),
        }
        if sha in seen_sha:
            for r in records:
                if r.get("source_sha256") == sha:
                    paths = set(r.get("all_paths", []))
                    paths.add(str(fpath))
                    r["all_paths"] = sorted(paths)
                    break
        else:
            records.append(rec)
            seen_sha.add(sha)
        # Print per-stage timing logs (reflect real success: whether any artifact exists)
        timing_info = out.get("timings", {})
        timing_str = " ".join([f"{k}={v:.2f}s" for k, v in timing_info.items() if isinstance(v, (int, float))])

        if rec["extract_ok"]:
            print(f"[OK] extracted {fpath}  → sha={sha[:8]} tool={rec['extract_tool']} ocr={rec['ocr_used']} timings: {timing_str}")
        else:
            # This includes unsupported binaries and parse failures (e.g., missing bs4), where we produced no sidecar.
            print(f"[MISS] no sidecar written for {fpath}  → sha={sha[:8]} tool={rec['extract_tool']} timings: {timing_str}")

    # Run verification if requested (or if verify-only)
    if args.verify or args.verify_only:
        df_for_verify = pd.DataFrame.from_records(records) if records else None
        run_verification(all_expected_paths, df_for_verify)

    if records:
        df_new = pd.DataFrame.from_records(records)
        out_path = parquet_dir / "sources.parquet"

        # By default, merge with existing manifest so extension-limited re-runs don't wipe earlier rows.
        df_out = df_new
        if (not args.no_merge_manifest) and out_path.exists():
            try:
                df_old = pd.read_parquet(out_path)

                # Concatenate then merge by sha, unioning all_paths and preferring non-null fields.
                df_cat = pd.concat([df_old, df_new], ignore_index=True)

                def _merge_group(g: pd.DataFrame) -> pd.Series:
                    # Prefer earliest first_seen_path for stability
                    first_seen = g["first_seen_path"].dropna().astype(str).iloc[0] if g["first_seen_path"].notna().any() else None

                    # Union all_paths (handle None/str/list/numpy)
                    paths = set()
                    for v in g.get("all_paths", pd.Series([], dtype=object)).dropna().tolist():
                        if v is None:
                            continue
                        if isinstance(v, str):
                            paths.add(v)
                        elif hasattr(v, "tolist"):
                            for p in v.tolist():
                                paths.add(str(p))
                        elif isinstance(v, (list, tuple, set)):
                            for p in v:
                                paths.add(str(p))
                        else:
                            paths.add(str(v))
                    all_paths = sorted(paths)

                    # For other columns, take the last non-null value (newer run wins), except extract_ok which is OR.
                    out = {}
                    for col in g.columns:
                        if col in ("all_paths", "first_seen_path"):
                            continue
                        if col == "extract_ok":
                            out[col] = bool(g["extract_ok"].fillna(False).any())
                            continue
                        # last non-null wins
                        nn = g[col].dropna()
                        out[col] = nn.iloc[-1] if len(nn) else None

                    out["first_seen_path"] = first_seen
                    out["all_paths"] = all_paths
                    return pd.Series(out)

                if "source_sha256" in df_cat.columns:
                    # Pandas is deprecating GroupBy.apply operating on grouping columns.
                    # Newer pandas supports include_groups=False; older versions do not.
                    gb = df_cat.groupby("source_sha256", dropna=False, sort=False)
                    try:
                        df_out = gb.apply(_merge_group, include_groups=False)
                    except TypeError:
                        # Fallback for older pandas
                        df_out = gb.apply(_merge_group)

                    # groupby+apply creates a multiindex with the group key; normalize to a clean frame.
                    df_out = df_out.reset_index(drop=False)
                    # Ensure we have a plain 'source_sha256' column.
                    if "source_sha256" not in df_out.columns:
                        # Some pandas versions name it 'level_0'
                        if "level_0" in df_out.columns:
                            df_out = df_out.rename(columns={"level_0": "source_sha256"})
                    # Drop any accidental index columns introduced by reset_index
                    for c in ["level_1", "index"]:
                        if c in df_out.columns:
                            df_out = df_out.drop(columns=[c])
                else:
                    df_out = df_cat

            except Exception as e:
                print(f"[WARN] Could not merge existing manifest {out_path}: {e}")
                df_out = df_new

        df_out.to_parquet(out_path, index=False)
        print(f"[OK] wrote manifest {out_path}  rows={len(df_out)}")
    else:
        print("[INFO] No supported files found during extraction.")